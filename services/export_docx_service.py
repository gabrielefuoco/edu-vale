import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database.connection import get_collection

async def export_diari_to_docx(user_id: str, utente: str = None, data_inizio: str = None, data_fine: str = None) -> str:
    """
    Esporta i diari di bordo in un file .docx.
    Filtri opzionali per nome utente e range di date.
    Ritorna il path assoluto del file generato.
    """
    col = await get_collection("diari_bordo", uid=user_id)
    
    query = {}
    if utente:
        pattern = re.compile(f'^{re.escape(utente)}$', re.IGNORECASE)
        query["utente"] = {"$regex": pattern}
        
    if data_inizio or data_fine:
        query["data"] = {}
        if data_inizio:
            query["data"]["$gte"] = data_inizio
        if data_fine:
            query["data"]["$lte"] = data_fine
            
    cursor = col.find(query).sort("data", -1).limit(500)
    diari = await cursor.to_list(length=500)
    
    if not diari:
        return None
        
    # Costruisci il nome file
    filename_parts = ["export_diari"]
    if utente:
        filename_parts.append(utente.replace(" ", "_"))
    if data_inizio:
        filename_parts.append(f"da_{data_inizio}")
    if data_fine:
        filename_parts.append(f"a_{data_fine}")
    filename = f"{'_'.join(filename_parts)}.docx"
    
    filepath = os.path.abspath(os.path.join(os.getcwd(), filename))
    
    # Crea il documento
    document = Document()
    
    # Non creiamo una pagina iniziale, iniziamo direttamente con i diari.
    
    # Inserisci i diari
    for d in diari:
        # Titolo diario (Data - Utente)
        data_diario = d.get('data', 'N/A')
        titolo_diario = document.add_heading(level=2)
        titolo_diario.add_run(f"Diario del {data_diario}").bold = True
        
        if not utente:
            p_user = document.add_paragraph("Utente: ")
            p_user.add_run(d.get('utente', 'Sconosciuto')).bold = True
            
        testo = d.get('testo_generato', '')
        
        # Inserisci il testo preservando i paragrafi
        for line in testo.split('\n'):
            line = line.strip()
            if not line:
                document.add_paragraph()
                continue
                
            # Salta la linea della data se già presente nel testo generato (per evitare duplicati)
            if line.lower().startswith('data:') or line.lower().startswith('**data:**'):
                continue
            
            # Gestione bullet point basica
            if line.startswith('- ') or line.startswith('* '):
                p = document.add_paragraph(line[2:], style='List Bullet')
            else:
                document.add_paragraph(line)
                
        # Spaziatura tra diari
        document.add_paragraph("-" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    document.save(filepath)
    return filepath
